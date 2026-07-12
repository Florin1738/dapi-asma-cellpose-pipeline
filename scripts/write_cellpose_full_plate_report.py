#!/usr/bin/env python
from __future__ import annotations

import argparse
import csv
from collections import Counter
from datetime import UTC, datetime
from pathlib import Path
from statistics import median
from typing import Any


FIGURE_FILES = [
    (
        "Figure 1",
        "figure_1_endpoint_metric_contrast.png",
        "Endpoint metric contrast across representative fields.",
    ),
    (
        "Figure 2",
        "figure_2_masking_effect_contrast.png",
        "Cellpose retained signal, retained area, and DAPI-positive nucleus denominator.",
    ),
    (
        "Figure 3",
        "figure_3_plate_level_endpoint_summary.png",
        "Plate-level endpoint summary.",
    ),
    (
        "Figure 4",
        "figure_4_representative_cell_segmentation_examples.png",
        "Representative Cellpose retained-region examples across low, middle, and high aSMA signal.",
    ),
]

CELLPOSE_2021_CITATION = (
    "Stringer C, Wang T, Michaelos M, Pachitariu M. PMID: 33318659."
)
CELLPOSE_2022_CITATION = (
    "Pachitariu M, Stringer C. PMID: 36344832."
)
CELLPOSE_DOCS_CITATION = (
    "Software documentation: Cellpose documentation, Cellpose-SAM and "
    "pretrained model documentation. PMID: not applicable. "
    "https://cellpose.readthedocs.io/"
)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Write Markdown, DOCX, and figure-caption PDFs for the full-plate Cellpose report."
    )
    parser.add_argument("--summary", required=True, type=Path)
    parser.add_argument("--figures-dir", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--workbook-path", required=True, type=Path)
    parser.add_argument("--run-root", required=True, type=Path)
    args = parser.parse_args()

    rows = _read_rows(args.summary)
    args.output_dir.mkdir(parents=True, exist_ok=True)
    stats = _summary_stats(rows)
    markdown = _build_markdown(
        stats=stats,
        figures_dir=args.figures_dir,
        workbook_path=args.workbook_path,
        run_root=args.run_root,
        summary_path=args.summary,
    )
    md_path = args.output_dir / "METHODS_AND_RESULTS.md"
    md_path.write_text(markdown, encoding="utf-8")

    docx_path = args.output_dir / "METHODS_AND_RESULTS.docx"
    pdf_dir = args.output_dir / "figure_pdfs"
    _write_docx(
        docx_path=docx_path,
        stats=stats,
        figures_dir=args.figures_dir,
        workbook_path=args.workbook_path,
        summary_path=args.summary,
    )
    _write_figure_pdfs(figures_dir=args.figures_dir, output_dir=pdf_dir)
    print(f"markdown={md_path}")
    print(f"docx={docx_path}")
    print(f"figure_pdfs={pdf_dir}")


def _read_rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def _summary_stats(rows: list[dict[str, str]]) -> dict[str, Any]:
    if not rows:
        raise ValueError("summary table is empty")
    plates = sorted({row["plate"] for row in rows}, key=_plate_sort_key)
    plate_counts = {plate: sum(row["plate"] == plate for row in rows) for plate in plates}
    whole_per_nuc = _values(rows, "whole_field_ch2_integrated_raw_per_DAPI_positive_nucleus")
    cellpose_per_nuc = _values(
        rows,
        "dapi_anchored_cellpose_ch2_integrated_background_corrected_per_DAPI_positive_nucleus",
    )
    cellpose_total = _values(rows, "dapi_anchored_cellpose_ch2_integrated_background_corrected")
    whole_total = _values(rows, "whole_field_ch2_integrated_raw")
    cellpose_area_per_nuc = _values(
        rows,
        "dapi_anchored_cellpose_masked_area_per_DAPI_positive_nucleus",
    )
    excluded_objects = _values(rows, "no_dapi_cellpose_object_count_excluded_in_anchored_variant")
    qc_status_counts = Counter(row.get("qc_status", "") or "missing_qc_status" for row in rows)
    fields_with_excluded_objects = sum(value > 0 for value in excluded_objects)
    retained_fraction = [
        retained / whole
        for retained, whole in zip(cellpose_total, whole_total, strict=True)
        if whole > 0
    ]
    top_rows = sorted(
        rows,
        key=lambda row: _number(
            row,
            "dapi_anchored_cellpose_ch2_integrated_background_corrected_per_DAPI_positive_nucleus",
        ),
        reverse=True,
    )[:5]
    return {
        "n_fields": len(rows),
        "plates": plates,
        "plate_counts": plate_counts,
        "whole_per_nucleus": _range_stats(whole_per_nuc),
        "cellpose_per_nucleus": _range_stats(cellpose_per_nuc),
        "cellpose_area_per_nucleus": _range_stats(cellpose_area_per_nuc),
        "retained_fraction": _range_stats(retained_fraction),
        "excluded_object_total": int(sum(excluded_objects)),
        "fields_with_excluded_objects": fields_with_excluded_objects,
        "qc_status_counts": dict(sorted(qc_status_counts.items())),
        "top_rows": top_rows,
    }


def _build_markdown(
    *,
    stats: dict[str, Any],
    figures_dir: Path,
    workbook_path: Path,
    run_root: Path,
    summary_path: Path,
) -> str:
    generated = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    plate_counts = "; ".join(
        f"{plate}: {count} fields" for plate, count in stats["plate_counts"].items()
    )
    top_fields = ", ".join(
        f"{row['source_id']} ({_sci(_number(row, 'dapi_anchored_cellpose_ch2_integrated_background_corrected_per_DAPI_positive_nucleus'))})"
        for row in stats["top_rows"]
    )
    figure_blocks = []
    for label, filename, caption in FIGURE_FILES:
        figure_blocks.extend(
            [
                f"![{label}. {caption}](figures/{filename})",
                "",
                f"**{label}. {caption}**",
                "",
            ]
        )
    return "\n".join(
        [
            "# Cellpose CH2+CH4 Quantification of alpha-smooth muscle actin Immunofluorescence",
            "",
            f"Generated: {generated}",
            "",
            "## Methods",
            "",
            "Microscopy fields were analyzed from TIFF exports organized by plate and acquisition folder. CH2 was treated as the alpha-smooth muscle actin (aSMA) fluorescence channel, and CH4 was treated as the DAPI nuclear counterstain channel. Images were read as primary two-dimensional TIFF planes; no Z projection was applied in this analysis.",
            "",
            "DAPI-positive nuclei were counted from the CH4 channel using the existing Cellpose nuclei masks generated for each acquisition folder. The denominator used throughout this report is the number of DAPI-positive nuclei, defined as the number of nonzero labels in the corresponding DAPI nucleus mask. DAPI fluorescence intensity itself was not used for normalization.",
            "",
            f"Candidate aSMA-associated regions were segmented with Cellpose using a two-channel input composed of CH2/aSMA and CH4/DAPI. The run used the pretrained `cpsam_v2` model through Cellpose. Cellpose is a generalist neural-network segmentation framework for cellular and nuclear image segmentation ({CELLPOSE_2021_CITATION}), and Cellpose-SAM/cpsam models are distributed in current Cellpose releases for broad pretrained segmentation use ({CELLPOSE_DOCS_CITATION}). In this workflow, the Cellpose output is treated as a candidate CH2+CH4 aSMA-associated region mask, not as validated whole-cell segmentation.",
            "",
            "For each image, the Cellpose retained-region mask was defined at the object level. A Cellpose object was retained only when at least one DAPI-positive nucleus centroid fell inside that object. Objects without a DAPI-positive nucleus centroid were excluded before intensity and area measurements were reported. Throughout this package, Cellpose refers to this DAPI-anchored retained-region mask.",
            "",
            "The Cellpose aSMA intensity was calculated as the sum of raw CH2 pixel intensities inside the retained Cellpose regions. No scalar per-pixel background subtraction was applied in this run (`background_value_per_px = 0`). Therefore, the difference between the whole-field measurement and the Cellpose measurement is produced by excluding CH2 pixels outside the retained Cellpose regions, not by subtracting a fitted background value from retained pixels.",
            "",
            "The primary per-nucleus intensity measurement was calculated as:",
            "",
            "`Cellpose retained-region CH2 integrated intensity / DAPI-positive nucleus count`",
            "",
            "This object-level DAPI-anchoring rule can remove separated no-DAPI candidate objects, but it does not solve every biological edge case, such as merged Cellpose objects containing multiple nuclei or true DAPI-negative cells.",
            "",
            "Automated QC status and QC flags from the Cellpose run were carried into the merged table and workbook. These flags are not manual validation. They indicate rows that need visual review or were rejected by predefined automated checks, and they should be considered before using row-level values downstream.",
            "",
            f"An optional extension, not applied here, would be to manually annotate a stratified subset of images and fine-tune Cellpose to this specific image domain. Cellpose 2.0 describes a human-in-the-loop annotation and model retraining workflow that can improve domain-specific segmentation when pretrained models are insufficient ({CELLPOSE_2022_CITATION}).",
            "",
            "## Results",
            "",
            f"The full-plate run analyzed {stats['n_fields']} fields across {len(stats['plates'])} plates ({plate_counts}).",
            "",
            f"Automated Cellpose candidate-region QC status was: {_format_status_counts(stats['qc_status_counts'])}. The descriptive summaries below include all processed fields; row-level QC status and flags are provided in the summary table and workbook for review.",
            "",
            f"The median whole-field raw CH2/aSMA intensity per DAPI-positive nucleus was {_sci(stats['whole_per_nucleus']['median'])}, with a range of {_sci(stats['whole_per_nucleus']['min'])} to {_sci(stats['whole_per_nucleus']['max'])}. After restricting the measurement to Cellpose retained regions, the median CH2/aSMA intensity per DAPI-positive nucleus was {_sci(stats['cellpose_per_nucleus']['median'])}, with a range of {_sci(stats['cellpose_per_nucleus']['min'])} to {_sci(stats['cellpose_per_nucleus']['max'])}.",
            "",
            f"Across fields, the median fraction of whole-field raw CH2/aSMA signal retained inside Cellpose regions was {_pct(stats['retained_fraction']['median'])}. This value describes how much CH2 signal remains after excluding pixels outside the retained Cellpose region.",
            "",
            f"The object-level DAPI-anchoring rule excluded {stats['excluded_object_total']} no-DAPI Cellpose objects across {stats['fields_with_excluded_objects']} fields before the Cellpose retained-region endpoint was reported.",
            "",
            f"The highest fields by Cellpose retained-region CH2/aSMA intensity per DAPI-positive nucleus were: {top_fields}.",
            "",
            "The positive-area endpoint is reported separately from integrated intensity. The median Cellpose retained area per DAPI-positive nucleus was "
            f"{stats['cellpose_area_per_nucleus']['median']:.1f} px/nucleus.",
            "",
            "## Figures",
            "",
            *figure_blocks,
            "Additional multi-page visual QC overlays are located in `figures/cellpose_overlay_pages/`. Each tile shows the raw CH2/aSMA intensity image in red on the left and the Cellpose retained-region overlay on the right. The overlay shows retained regions in transparent green, DAPI signal/centroid marks in blue, and orange boundaries for Cellpose objects that contain no DAPI nucleus centroid. `figures/cellpose_overlay_pages/overlay_index.csv` maps each `source_id` to the overlay page and tile number.",
            "",
            "## Outputs",
            "",
            f"- Machine-readable summary table: `{summary_path}`",
            f"- Excel workbook: `{workbook_path}`",
            f"- Full Cellpose run root: `{run_root}`",
            "",
            "## References",
            "",
            CELLPOSE_2021_CITATION,
            "",
            CELLPOSE_2022_CITATION,
            "",
            CELLPOSE_DOCS_CITATION,
            "",
        ]
    )


def _write_docx(
    *,
    docx_path: Path,
    stats: dict[str, Any],
    figures_dir: Path,
    workbook_path: Path,
    summary_path: Path,
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run("Cellpose CH2+CH4 Quantification of alpha-smooth muscle actin Immunofluorescence")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _rgb("17324D")
    title.paragraph_format.space_after = Pt(10)

    _add_key_value_table(
        doc,
        [
            ("Fields analyzed", str(stats["n_fields"])),
            ("Plates", "; ".join(f"{k}: {v}" for k, v in stats["plate_counts"].items())),
            ("Automated QC statuses", _format_status_counts(stats["qc_status_counts"])),
            ("Workbook", "PI-style retained-region workbook"),
            ("Summary CSV", "full-plate endpoint summary CSV"),
        ],
    )
    _heading(doc, "Methods", level=1)
    for paragraph in _method_paragraphs():
        doc.add_paragraph(paragraph)
    _heading(doc, "Results", level=1)
    for paragraph in _result_paragraphs(stats):
        doc.add_paragraph(paragraph)
    _heading(doc, "Figures", level=1)
    for label, filename, caption in FIGURE_FILES:
        figure_path = figures_dir / filename
        if figure_path.exists():
            doc.add_picture(str(figure_path), width=Inches(6.3))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap_run = cap.add_run(f"{label}. {caption}")
        cap_run.bold = True
        cap.paragraph_format.space_after = Pt(8)
    _heading(doc, "References", level=1)
    for reference in [
        CELLPOSE_2021_CITATION,
        CELLPOSE_2022_CITATION,
        CELLPOSE_DOCS_CITATION,
    ]:
        doc.add_paragraph(reference)
    doc.save(docx_path)


def _write_figure_pdfs(*, figures_dir: Path, output_dir: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    from PIL import Image

    output_dir.mkdir(parents=True, exist_ok=True)
    for label, filename, caption in FIGURE_FILES:
        figure_path = figures_dir / filename
        if not figure_path.exists():
            continue
        pdf_path = output_dir / f"{Path(filename).stem}.pdf"
        image = Image.open(figure_path)
        page_w, page_h = letter
        margin = 0.45 * inch
        caption_h = 0.8 * inch
        available_w = page_w - 2 * margin
        available_h = page_h - 2 * margin - caption_h
        scale = min(available_w / image.width, available_h / image.height)
        draw_w = image.width * scale
        draw_h = image.height * scale
        pdf = canvas.Canvas(str(pdf_path), pagesize=letter)
        pdf.drawImage(
            str(figure_path),
            (page_w - draw_w) / 2,
            margin + caption_h,
            width=draw_w,
            height=draw_h,
            preserveAspectRatio=True,
            anchor="c",
        )
        pdf.setFont("Helvetica-Bold", 10)
        text = pdf.beginText(margin, margin + 0.45 * inch)
        for line in _wrap_text(f"{label}. {caption}", 95):
            text.textLine(line)
        pdf.drawText(text)
        pdf.showPage()
        pdf.save()


def _method_paragraphs() -> list[str]:
    return [
        "Microscopy fields were analyzed from TIFF exports organized by plate and acquisition folder. CH2 was treated as the alpha-smooth muscle actin (aSMA) fluorescence channel, and CH4 was treated as the DAPI nuclear counterstain channel. Images were read as primary two-dimensional TIFF planes; no Z projection was applied in this analysis.",
        "DAPI-positive nuclei were counted from the CH4 channel using the existing Cellpose nuclei masks generated for each acquisition folder. The denominator used throughout this report is the number of DAPI-positive nuclei, defined as the number of nonzero labels in the corresponding DAPI nucleus mask. DAPI fluorescence intensity itself was not used for normalization.",
        f"Candidate aSMA-associated regions were segmented with Cellpose using a two-channel input composed of CH2/aSMA and CH4/DAPI. The run used the pretrained cpsam_v2 model through Cellpose. Cellpose is a generalist neural-network segmentation framework for cellular and nuclear image segmentation ({CELLPOSE_2021_CITATION}), and Cellpose-SAM/cpsam models are distributed in current Cellpose releases for broad pretrained segmentation use ({CELLPOSE_DOCS_CITATION}). In this workflow, the Cellpose output is treated as a candidate CH2+CH4 aSMA-associated region mask, not as validated whole-cell segmentation.",
        "For each image, the Cellpose retained-region mask was defined at the object level. A Cellpose object was retained only when at least one DAPI-positive nucleus centroid fell inside that object. Objects without a DAPI-positive nucleus centroid were excluded before intensity and area measurements were reported. Throughout this package, Cellpose refers to this DAPI-anchored retained-region mask.",
        "The Cellpose aSMA intensity was calculated as the sum of raw CH2 pixel intensities inside the retained Cellpose regions. No scalar per-pixel background subtraction was applied in this run (background_value_per_px = 0). Therefore, the difference between the whole-field measurement and the Cellpose measurement is produced by excluding CH2 pixels outside the retained Cellpose regions, not by subtracting a fitted background value from retained pixels.",
        "The primary per-nucleus intensity measurement was calculated as Cellpose retained-region CH2 integrated intensity divided by DAPI-positive nucleus count.",
        "The object-level DAPI-anchoring rule can remove separated no-DAPI candidate objects, but it does not solve every biological edge case, such as merged Cellpose objects containing multiple nuclei or true DAPI-negative cells.",
        "Automated QC status and QC flags from the Cellpose run were carried into the merged table and workbook. These flags are not manual validation. They indicate rows that need visual review or were rejected by predefined automated checks, and they should be considered before using row-level values downstream.",
        f"An optional extension, not applied here, would be to manually annotate a stratified subset of images and fine-tune Cellpose to this specific image domain. Cellpose 2.0 describes a human-in-the-loop annotation and model retraining workflow that can improve domain-specific segmentation when pretrained models are insufficient ({CELLPOSE_2022_CITATION}).",
    ]


def _result_paragraphs(stats: dict[str, Any]) -> list[str]:
    plate_counts = "; ".join(f"{plate}: {count}" for plate, count in stats["plate_counts"].items())
    return [
        f"The run analyzed {stats['n_fields']} fields across {len(stats['plates'])} plates ({plate_counts}).",
        f"Automated Cellpose candidate-region QC status was: {_format_status_counts(stats['qc_status_counts'])}. The descriptive summaries in this report include all processed fields; row-level QC status and flags are provided in the summary table and workbook for review.",
        f"The median whole-field raw CH2/aSMA intensity per DAPI-positive nucleus was {_sci(stats['whole_per_nucleus']['median'])}, with a range of {_sci(stats['whole_per_nucleus']['min'])} to {_sci(stats['whole_per_nucleus']['max'])}. After restricting the measurement to Cellpose retained regions, the median CH2/aSMA intensity per DAPI-positive nucleus was {_sci(stats['cellpose_per_nucleus']['median'])}, with a range of {_sci(stats['cellpose_per_nucleus']['min'])} to {_sci(stats['cellpose_per_nucleus']['max'])}.",
        f"Across fields, the median fraction of whole-field raw CH2/aSMA signal retained inside Cellpose regions was {_pct(stats['retained_fraction']['median'])}. This value describes how much CH2 signal remains after excluding pixels outside the retained Cellpose region.",
        f"The object-level DAPI-anchoring rule excluded {stats['excluded_object_total']} no-DAPI Cellpose objects across {stats['fields_with_excluded_objects']} fields before the Cellpose retained-region endpoint was reported.",
        "The positive-area endpoint is reported separately from integrated intensity. The median Cellpose retained area per DAPI-positive nucleus was "
        f"{stats['cellpose_area_per_nucleus']['median']:.1f} px/nucleus.",
    ]


def _add_key_value_table(doc: Any, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for (label, value), cells in zip(rows, table.rows, strict=True):
        cells.cells[0].text = label
        cells.cells[1].text = value
        for run in cells.cells[0].paragraphs[0].runs:
            run.bold = True


def _heading(doc: Any, text: str, *, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = _rgb("2E74B5")


def _values(rows: list[dict[str, str]], column: str) -> list[float]:
    return [_number(row, column) for row in rows if _number(row, column) == _number(row, column)]


def _number(row: dict[str, str], column: str) -> float:
    try:
        return float(row[column])
    except (KeyError, TypeError, ValueError):
        return float("nan")


def _range_stats(values: list[float]) -> dict[str, float]:
    clean = [value for value in values if value == value]
    return {
        "min": min(clean),
        "median": median(clean),
        "max": max(clean),
    }


def _sci(value: float) -> str:
    if value != value:
        return "NA"
    return f"{value:.2e}"


def _pct(value: float) -> str:
    if value != value:
        return "NA"
    return f"{100.0 * value:.1f}%"


def _format_status_counts(counts: dict[str, int]) -> str:
    order = ["reviewable_not_validated", "needs_manual_review", "reject_qc_failure", "missing_qc_status"]
    items = [(key, counts[key]) for key in order if key in counts]
    items.extend((key, value) for key, value in counts.items() if key not in order)
    return ", ".join(f"{key}: {value}" for key, value in items)


def _plate_sort_key(value: str) -> tuple[int, str]:
    digits = "".join(character for character in value if character.isdigit())
    return (int(digits) if digits else 999, value)


def _rgb(hex_value: str) -> Any:
    from docx.shared import RGBColor

    value = hex_value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _wrap_text(text: str, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        if sum(len(item) + 1 for item in current) + len(word) > width:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


if __name__ == "__main__":
    main()
